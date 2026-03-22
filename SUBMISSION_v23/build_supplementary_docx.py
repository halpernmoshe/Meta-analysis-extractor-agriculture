#!/usr/bin/env python3
"""
Build supplementary materials DOCX from CSVs and PNGs.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

import csv
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = Path(__file__).parent
SUPP_DIR = SCRIPT_DIR / "supplementary"
OUTPUT_FILE = SCRIPT_DIR / "halpern_2026_supplementary.docx"

# Table definitions: (csv_filename, caption)
TABLES = [
    (
        "Table_S1_TOST_results.csv",
        "Table S1. TOST equivalence results for all five datasets at four margin levels "
        "(\u00b12 pp, \u00b13 pp, proportional \u00b120%, and proportional \u00b110% of mean "
        "absolute effect size). CR2 bias-corrected sandwich estimator with Satterthwaite "
        "degrees of freedom."
    ),
    (
        "Table_S2_per_paper_agreement.csv",
        "Table S2. Per-paper agreement statistics for all papers across five datasets. "
        "Tier classification: Excellent (MAE < 5 pp), Good (5\u201310 pp), Fair (10\u201320 pp), "
        "Poor (> 20 pp)."
    ),
    (
        "Table_S3_variance_recovery.csv",
        "Table S3. Variance recovery details by dataset, including direct extraction coverage, "
        "indirect recovery, and imputation sensitivity."
    ),
    (
        "Table_S4_agent_replication.csv",
        "Table S4. Agent replication stability. Aggregate effect difference between independent "
        "duplicate runs."
    ),
]

# Figure definitions: (png_filename, caption)
FIGURES = [
    (
        "Figure_S1_per_element_effects.png",
        "Figure S1. Per-element effect sizes for the Loladze 2014 dataset. Agent-extracted "
        "(orange) versus reference (blue) mean percentage change under elevated CO\u2082. "
        "Elements marked with * (Fe, Mn) increase under elevated CO\u2082, which is biologically "
        "correct. Error reflects alignment and extraction combined."
    ),
    (
        "Figure_S2_source_type_distribution.png",
        "Figure S2. Source-type distribution across datasets. The Li 2024 (biochar) dataset "
        "has detailed source labeling; other datasets show available classification."
    ),
    (
        "Figure_S3_variance_sensitivity.png",
        "Figure S3. Variance recovery sensitivity analysis for the Li 2024 (biochar) dataset. "
        "Pooled effect size under five imputation strategies, demonstrating robustness "
        "(spread < 1 pp)."
    ),
]


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def read_csv(path):
    with open(path, "r", encoding="utf-8") as f:
        reader = csv.reader(f)
        return list(reader)


def add_table(doc, rows):
    if not rows:
        return
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j >= num_cols:
                break
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT

            run = p.add_run(cell_text)
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"

            if i == 0:
                run.bold = True
                set_cell_shading(cell, "2F5496")
                run.font.color.rgb = RGBColor(255, 255, 255)
            elif i % 2 == 0:
                set_cell_shading(cell, "D6E4F0")


def build():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # --- Title ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Supplementary Materials")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    # --- Subtitle ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(
        "Breaking the Extraction Bottleneck: A Single AI Agent Achieves Statistical "
        "Equivalence with Human-Extracted Meta-Analysis Data Across Five Agricultural Datasets"
    )
    run.italic = True
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # --- Author ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("Moshe Halpern")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"

    # --- Tables ---
    for csv_file, caption in TABLES:
        csv_path = SUPP_DIR / csv_file
        if not csv_path.exists():
            print(f"WARNING: {csv_path} not found, skipping")
            continue

        # Caption
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(caption)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"

        # Table
        rows = read_csv(csv_path)
        add_table(doc, rows)

        # Spacing
        doc.add_paragraph()

    # --- Figures ---
    for png_file, caption in FIGURES:
        png_path = SUPP_DIR / png_file
        if not png_path.exists():
            print(f"WARNING: {png_path} not found, skipping")
            continue

        # Page break
        doc.add_page_break()

        # Image
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(png_path), width=Inches(6.0))

        # Caption
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run(caption)
        run.italic = True
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"

    # Save
    doc.save(str(OUTPUT_FILE))
    size = OUTPUT_FILE.stat().st_size
    print(f"Saved: {OUTPUT_FILE}")
    print(f"Size: {size:,} bytes ({size/1024:.1f} KB)")


if __name__ == "__main__":
    build()
