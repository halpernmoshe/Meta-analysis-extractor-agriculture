#!/usr/bin/env python3
"""
Build a publication-ready DOCX from PAPER_FINAL_v19.md with embedded figures.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = Path(__file__).parent
MD_FILE = SCRIPT_DIR / "PAPER_FINAL_v23.md"
FIGURES_DIR = SCRIPT_DIR / "figures"
OUTPUT_FILE = SCRIPT_DIR.parent / "halpern_2026_v3.docx"

# Map figure numbers to files
FIGURE_FILES = {
    1: "figure2_scatter_plots.png",       # Scatter plots (referenced as FIGURE 1)
    2: "figure3_cross_dataset_comparison.png",  # Per-paper MAE (referenced as FIGURE 2)
    3: "figure4_bland_altman.png",         # Bland-Altman (referenced as FIGURE 3)
    4: "figure1_architecture.png",         # Cross-method (referenced as FIGURE 4)
    5: "figure5_source_type_accuracy.png", # Source type (referenced as FIGURE 5)
    6: "figure6_aggregate_effects.png",    # Aggregate effects
}


def set_cell_shading(cell, color_hex):
    """Set background shading for a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_page_number(doc):
    """Add page numbers to footer."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add page number field
    run = p.add_run()
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fld_char_begin)

    run2 = p.add_run()
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instr)

    run3 = p.add_run()
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fld_char_end)


def add_formatted_run(paragraph, text, bold=False, italic=False, font_size=None, font_name=None):
    """Add a run with formatting to a paragraph."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    if font_name:
        run.font.name = font_name
    return run


def _add_run_with_superscripts(paragraph, text, base_size=12, base_font='Times New Roman',
                                bold=False, italic=False):
    """Add text to paragraph, converting ^...^ to superscript runs."""
    # Split on superscript markers ^...^
    sup_parts = re.split(r'(\^[^^]+?\^)', text)
    for sp in sup_parts:
        if not sp:
            continue
        if sp.startswith('^') and sp.endswith('^'):
            inner = sp[1:-1]
            run = paragraph.add_run(inner)
            run.font.superscript = True
            run.font.size = Pt(max(base_size - 2, 8))
            run.font.name = base_font
            run.bold = bold
            run.italic = italic
        else:
            clean = sp.replace('---', '\u2014').replace('--', '\u2013')
            run = paragraph.add_run(clean)
            run.font.size = Pt(base_size)
            run.font.name = base_font
            run.bold = bold
            run.italic = italic


def parse_inline_formatting(paragraph, text, base_size=12, base_font='Times New Roman'):
    """Parse markdown inline formatting (**bold**, *italic*, ^superscript^) into Word runs."""
    # Pattern to match **bold**, *italic*, and plain text
    # Process bold first, then italic
    parts = re.split(r'(\*\*.*?\*\*|\*[^*]+?\*)', text)

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            inner = part[2:-2]
            _add_run_with_superscripts(paragraph, inner, base_size, base_font, bold=True)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            inner = part[1:-1]
            _add_run_with_superscripts(paragraph, inner, base_size, base_font, italic=True)
        else:
            _add_run_with_superscripts(paragraph, part, base_size, base_font)


def insert_figure(doc, fig_num, caption_text):
    """Insert a figure image and caption."""
    # Try to find the figure file
    fig_file = FIGURE_FILES.get(fig_num)
    fig_path = FIGURES_DIR / fig_file if fig_file else None

    if fig_path and fig_path.exists():
        # Add image paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(fig_path), width=Inches(6.0))

        # Add caption
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cap.paragraph_format.space_after = Pt(12)
        run_label = cap.add_run(f"Figure {fig_num}. ")
        run_label.bold = True
        run_label.font.size = Pt(10)
        run_label.font.name = 'Times New Roman'

        # Parse the rest of caption for inline formatting
        parse_inline_formatting(cap, caption_text, base_size=10)
    else:
        # Figure missing - add note
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[Figure {fig_num} image not found: {fig_file}]")
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(180, 0, 0)

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_label = cap.add_run(f"Figure {fig_num}. ")
        run_label.bold = True
        run_label.font.size = Pt(10)
        parse_inline_formatting(cap, caption_text, base_size=10)


def parse_table(lines):
    """Parse markdown table lines into list of rows (each row is list of cells)."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith('|'):
            continue
        # Skip separator lines
        if re.match(r'^\|[\s\-:|]+\|$', line):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]  # Remove first/last empty
        rows.append(cells)
    return rows


def add_word_table(doc, rows):
    """Add a formatted Word table with alternating row shading."""
    if not rows:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set table style
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j >= num_cols:
                break
            cell = row.cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT

            # First column left-aligned always
            if j == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Parse bold/italic/superscript in cell text
            clean_text = cell_text.replace('---', '\u2014').replace('--', '\u2013')
            if '**' in clean_text or '*' in clean_text or '^' in clean_text:
                parse_inline_formatting(p, clean_text, base_size=9, base_font='Times New Roman')
            else:
                run = p.add_run(clean_text)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'

            # Header row styling
            if i == 0:
                for run in p.runs:
                    run.bold = True
                set_cell_shading(cell, "2F5496")
                for run in p.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            # Alternating row shading
            elif i % 2 == 0:
                set_cell_shading(cell, "D6E4F0")

    # Add spacing after table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def build_docx():
    """Main function to build the DOCX."""
    print("Reading markdown file...")
    with open(MD_FILE, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')

    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # --- Default style ---
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    # Heading styles
    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Times New Roman'
        h_style.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            h_style.font.size = Pt(14)
        elif level == 2:
            h_style.font.size = Pt(13)
        else:
            h_style.font.size = Pt(12)

    # --- Process lines ---
    i = 0
    title_done = False
    in_references = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines (but add paragraph spacing naturally)
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped == '---':
            i += 1
            continue

        # Figure references [FIGURE N: caption]
        fig_match = re.match(r'^\[FIGURE\s+(\d+)[:\s]*(.*)\]$', stripped)
        if fig_match:
            fig_num = int(fig_match.group(1))
            caption = fig_match.group(2).strip()
            insert_figure(doc, fig_num, caption)
            i += 1
            continue

        # Tables: collect consecutive lines starting with |
        if stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            rows = parse_table(table_lines)
            if rows:
                add_word_table(doc, rows)
            continue

        # Table caption (bold line starting with **Table)
        table_cap_match = re.match(r'^\*\*Table\s+\d+\..*\*\*$', stripped)
        if table_cap_match:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            text = stripped.strip('*')
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()

            # Title (first H1)
            if level == 1 and not title_done:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(24)
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(heading_text)
                run.bold = True
                run.font.size = Pt(16)
                run.font.name = 'Times New Roman'
                title_done = True
                i += 1
                continue

            # Track if we're in references
            if 'References' in heading_text or 'Bibliography' in heading_text:
                in_references = True

            # Section headings
            h = doc.add_heading(heading_text, level=min(level, 3))
            # Remove numbering artifacts
            i += 1
            continue

        # Author line
        if stripped.startswith('**') and 'Halpern' in stripped and not title_done:
            # Already handled by title
            i += 1
            continue
        if 'Halpern' in stripped and '**' in stripped:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            name = stripped.replace('**', '').replace('^ORCID^', '').strip()
            run = p.add_run(name)
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            i += 1
            continue

        # Affiliation/correspondence
        if stripped.startswith('Institute of') or stripped.startswith('*Correspondence'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            clean = stripped.replace('*Correspondence:*', 'Correspondence:')
            parse_inline_formatting(p, clean, base_size=10)
            i += 1
            continue

        # Keywords line
        if stripped.startswith('**Keywords'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            parse_inline_formatting(p, stripped, base_size=11)
            i += 1
            continue

        # Bullet points
        if stripped.startswith('- ') or stripped.startswith('* '):
            bullet_text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            parse_inline_formatting(p, bullet_text, base_size=12)
            i += 1
            continue

        # Numbered list (but not in references - keep AMA numbers visible)
        num_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if num_match and not in_references:
            p = doc.add_paragraph(style='List Number')
            parse_inline_formatting(p, num_match.group(2), base_size=12)
            i += 1
            continue

        # Table note: italic line immediately after a table (starts and ends with *)
        if re.match(r'^\*[^*].*[^*]\*$', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            inner = stripped[1:-1]  # Remove surrounding *
            run = p.add_run(inner)
            run.italic = True
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            i += 1
            continue

        # Regular paragraph (possibly with references formatting)
        p = doc.add_paragraph()
        if in_references:
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.left_indent = Inches(0.5)
            parse_inline_formatting(p, stripped, base_size=11)
        else:
            parse_inline_formatting(p, stripped, base_size=12)

        i += 1

    # Add page numbers
    add_page_number(doc)

    # Save
    print(f"Saving to {OUTPUT_FILE}...")
    doc.save(str(OUTPUT_FILE))

    file_size = OUTPUT_FILE.stat().st_size
    print(f"Done! File size: {file_size:,} bytes ({file_size/1024:.1f} KB)")
    print(f"Output: {OUTPUT_FILE}")


if __name__ == '__main__':
    build_docx()
