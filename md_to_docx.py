"""Convert PAPER_FULL_DRAFT.md to a Word document with embedded figures."""
import sys, re, os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')

BASE = Path(r"C:\Users\moshe\Dropbox\Testing metaanalyis program\meta_analysis_extractor")
MD_FILE = BASE / "PAPER_FULL_DRAFT.md"
FIG_DIR = BASE / "output" / "paper_figures"
FORMAL_DIR = BASE / "output" / "formal_stats"
SENS_DIR = BASE / "output" / "sensitivity"
OUT_FILE = BASE / "PAPER_FULL_DRAFT.docx"

# Map figure references to actual files
FIGURE_MAP = {
    'Figure 1':  (FIG_DIR / "fig1_pipeline_architecture.png", "Figure 1. Pipeline architecture: four-stage workflow comprising challenge-aware reconnaissance, dual-model extraction (Claude Sonnet 4 + Kimi K2.5), consensus building with Gemini 3 Flash tiebreaker, and post-processing."),
    'Figure 2':  (FIG_DIR / "fig_combined_scatter.png", "Figure 2. Combined scatter plots across all three validation datasets: (A) Loladze 2014 (r = 0.669), (B) Hui 2023 (r = 0.950), (C) Li 2022 (r = 0.453). Points are colored by mineral element (A) or study (B, C). Dashed lines represent perfect agreement (y = x)."),
    'Figure 3':  (FIG_DIR / "fig_error_distribution.png", "Figure 3. Error distribution: histogram of absolute errors and cumulative percentage curve showing within-5% (58%), within-10% (74%), and within-20% (91%) thresholds."),
    'Figure 4':  (SENS_DIR / "fig_loo_combined.png", "Figure 4. Leave-one-out sensitivity analysis: (A) Leave-one-paper-out showing change in MAE when each paper is excluded, (B) Leave-one-element-out showing analogous results by element."),
    'Figure 5':  (FIG_DIR / "fig_tost_equivalence.png", "Figure 5. TOST equivalence forest plot across all three validation datasets with summary statistics."),
    'Figure 6':  (FIG_DIR / "fig_bland_altman_trio.png", "Figure 6. Bland-Altman analysis across all three datasets showing limits of agreement scaling with data complexity."),
}

# Where to insert each figure (after which section heading or keyword)
FIGURE_PLACEMENT = {
    'Figure 1':  '2.1 Pipeline Architecture',
    'Figure 2':  '3.3 Cross-Dataset Validation',
    'Figure 3':  '3.4 Extraction Method',
    'Figure 4':  '3.9.1 Leave-One-Paper',
    'Figure 5':  '3.8.4 Cross-Dataset Formal',
    'Figure 6':  '3.8.4 Cross-Dataset Formal',
}


def add_figure(doc, fig_key):
    """Insert a figure with caption."""
    fig_path, caption = FIGURE_MAP[fig_key]
    if not fig_path.exists():
        p = doc.add_paragraph()
        run = p.add_run(f"[{fig_key}: file not found at {fig_path.name}]")
        run.italic = True
        run.font.color.rgb = RGBColor(180, 0, 0)
        return

    # Add figure
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(fig_path), width=Inches(6.0))

    # Add caption
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap_p.paragraph_format.space_after = Pt(12)
    # Bold the "Figure N." part
    parts = caption.split('.', 1)
    run_bold = cap_p.add_run(parts[0] + '.')
    run_bold.bold = True
    run_bold.font.size = Pt(10)
    if len(parts) > 1:
        run_rest = cap_p.add_run(parts[1])
        run_rest.font.size = Pt(10)


def parse_table(lines):
    """Parse markdown table lines into rows of cells."""
    rows = []
    for line in lines:
        line = line.strip().strip('|')
        cells = [c.strip() for c in line.split('|')]
        rows.append(cells)
    # Remove separator row (----)
    rows = [r for r in rows if not all(set(c) <= set('-: ') for c in r)]
    return rows


def add_md_table(doc, md_rows):
    """Add a markdown table to the Word document."""
    if not md_rows or len(md_rows) < 2:
        return
    n_cols = len(md_rows[0])
    table = doc.add_table(rows=len(md_rows), cols=n_cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(md_rows):
        for j, cell_text in enumerate(row_data):
            if j < n_cols:
                cell = table.cell(i, j)
                # Clear default text and use inline formatting
                cell.text = ''
                paragraph = cell.paragraphs[0]
                apply_inline_formatting(paragraph, cell_text)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.space_before = Pt(2)
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                    if i == 0:
                        for run in paragraph.runs:
                            run.bold = True

    doc.add_paragraph()  # spacer


def apply_inline_formatting(paragraph, text):
    """Apply bold and italic markdown formatting to a paragraph."""
    # Split on **bold** and *italic* markers
    pattern = r'(\*\*.*?\*\*|\*.*?\*)'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def restart_numbering(paragraph):
    """Force a numbered list paragraph to restart at 1.

    python-docx's 'List Number' style relies on the style definition for numPr,
    so we must explicitly add numPr to the paragraph to create a new numbering
    instance with a startOverride.
    """
    # Get the numbering definitions from the document
    numbering_part = paragraph.part.numbering_part
    numbering = numbering_part.numbering_definitions._numbering

    # Find the numId used by the 'List Number' style
    # Look through existing num definitions to find one with the right abstract
    all_nums = numbering.findall(qn('w:num'))
    if not all_nums:
        return

    # Find an existing 'List Number' numId by checking the style's numbering
    style_elem = paragraph.style.element
    style_pPr = style_elem.find(qn('w:pPr'))
    style_numPr = style_pPr.find(qn('w:numPr')) if style_pPr is not None else None

    if style_numPr is None:
        return

    style_numId_elem = style_numPr.find(qn('w:numId'))
    if style_numId_elem is None:
        return

    style_numId = style_numId_elem.get(qn('w:val'))

    # Find the abstract numbering referenced by this numId
    absId = None
    for n in all_nums:
        if n.get(qn('w:numId')) == style_numId:
            absRef = n.find(qn('w:abstractNumId'))
            if absRef is not None:
                absId = absRef.get(qn('w:val'))
            break

    if absId is None:
        return

    # Create a new num that references the same abstract but with startOverride=1
    new_numId = str(max(int(n.get(qn('w:numId'))) for n in all_nums) + 1)

    new_num = OxmlElement('w:num')
    new_num.set(qn('w:numId'), new_numId)

    new_absRef = OxmlElement('w:abstractNumId')
    new_absRef.set(qn('w:val'), absId)
    new_num.append(new_absRef)

    lvl_override = OxmlElement('w:lvlOverride')
    lvl_override.set(qn('w:ilvl'), '0')
    start_override = OxmlElement('w:startOverride')
    start_override.set(qn('w:val'), '1')
    lvl_override.append(start_override)
    new_num.append(lvl_override)

    numbering.append(new_num)

    # Now add explicit numPr to this paragraph pointing to the new numId
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    numPr.append(ilvl)
    numId_elem = OxmlElement('w:numId')
    numId_elem.set(qn('w:val'), new_numId)
    numPr.append(numId_elem)
    pPr.append(numPr)


def process_markdown(doc, md_text):
    """Process markdown text and build the Word document."""
    lines = md_text.split('\n')
    i = 0
    current_section = ""
    figures_placed = set()
    pending_figures = {}  # section_key -> list of figure keys
    in_numbered_list = False  # Track whether previous line was a numbered list

    # Build reverse map: which figures to place after which sections
    for fig_key, section_hint in FIGURE_PLACEMENT.items():
        if section_hint not in pending_figures:
            pending_figures[section_hint] = []
        pending_figures[section_hint].append(fig_key)

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip horizontal rules
        if stripped == '---':
            in_numbered_list = False
            i += 1
            continue

        # Skip the "Working Draft" line
        if stripped.startswith('**Working Draft'):
            in_numbered_list = False
            i += 1
            continue

        # Headings
        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            heading_text = stripped.lstrip('#').strip()

            if level == 1:
                h = doc.add_heading(heading_text, level=1)
            elif level == 2:
                h = doc.add_heading(heading_text, level=2)
            elif level == 3:
                h = doc.add_heading(heading_text, level=3)
            else:
                h = doc.add_heading(heading_text, level=min(level, 4))

            current_section = heading_text
            in_numbered_list = False

            # Check if we should place figures after this section heading
            for section_hint, fig_keys in pending_figures.items():
                if section_hint.lower() in current_section.lower():
                    # Don't place immediately - place after the section content
                    pass

            i += 1
            continue

        # Table detection
        if '|' in stripped and i + 1 < len(lines) and '|' in lines[i + 1]:
            in_numbered_list = False
            table_lines = []
            while i < len(lines) and '|' in lines[i].strip():
                table_lines.append(lines[i])
                i += 1
            md_rows = parse_table(table_lines)
            if md_rows:
                add_md_table(doc, md_rows)
            continue

        # Bullet points
        if stripped.startswith('- ') or stripped.startswith('* '):
            in_numbered_list = False
            bullet_text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            apply_inline_formatting(p, bullet_text)
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if num_match:
            list_text = num_match.group(2)
            p = doc.add_paragraph(style='List Number')
            apply_inline_formatting(p, list_text)
            # If this is the first item in a new list, restart numbering
            if not in_numbered_list:
                restart_numbering(p)
            in_numbered_list = True
            i += 1
            continue

        # Empty lines
        if not stripped:
            in_numbered_list = False
            i += 1

            # Check if we should place figures before the next section
            # Look ahead to see if next non-empty line is a heading
            j = i
            while j < len(lines) and not lines[j].strip():
                j += 1

            if j < len(lines) and lines[j].strip().startswith('#'):
                # Before the next section, place any pending figures for current section
                for section_hint, fig_keys in pending_figures.items():
                    if section_hint.lower() in current_section.lower():
                        for fk in fig_keys:
                            if fk not in figures_placed:
                                add_figure(doc, fk)
                                figures_placed.add(fk)

            continue

        # Regular paragraphs
        in_numbered_list = False
        p = doc.add_paragraph()
        apply_inline_formatting(p, stripped)

        # Check if this paragraph references a figure that should be placed inline
        for fig_key in FIGURE_MAP:
            fig_ref = f"({fig_key})" if fig_key != 'Figure 1' else fig_key
            if fig_key in stripped and fig_key not in figures_placed:
                # Check if section matches placement
                for section_hint, fig_keys in pending_figures.items():
                    if fig_key in fig_keys and section_hint.lower() in current_section.lower():
                        add_figure(doc, fig_key)
                        figures_placed.add(fig_key)
                        break

        i += 1

    # Place any remaining figures at the end
    for fig_key in FIGURE_MAP:
        if fig_key not in figures_placed:
            add_figure(doc, fig_key)
            figures_placed.add(fig_key)


def main():
    print("Creating Word document from paper draft...")

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Set heading styles
    for level in range(1, 5):
        h_style = doc.styles[f'Heading {level}']
        h_font = h_style.font
        h_font.name = 'Times New Roman'
        if level == 1:
            h_font.size = Pt(16)
            h_font.bold = True
        elif level == 2:
            h_font.size = Pt(14)
            h_font.bold = True
        elif level == 3:
            h_font.size = Pt(12)
            h_font.bold = True
        else:
            h_font.size = Pt(11)
            h_font.bold = True

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Read markdown
    md_text = MD_FILE.read_text(encoding='utf-8')

    # Process
    process_markdown(doc, md_text)

    # Save
    doc.save(str(OUT_FILE))
    print(f"Saved: {OUT_FILE}")
    print(f"Size: {OUT_FILE.stat().st_size / 1024:.0f} KB")

    # Count figures placed
    figs_found = sum(1 for _, (p, _) in FIGURE_MAP.items() if p.exists())
    print(f"Figures embedded: {figs_found}/{len(FIGURE_MAP)}")


if __name__ == '__main__':
    main()
